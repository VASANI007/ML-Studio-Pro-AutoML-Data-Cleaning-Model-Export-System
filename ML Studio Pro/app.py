import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import sys, os, io, time, traceback

sys.path.insert(0, os.path.dirname(__file__))

st.set_page_config(page_title="ML Studio Pro", page_icon="⚡", layout="wide",
                   initial_sidebar_state="expanded")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Mono:wght@400;700&family=Syne:wght@400;600;700;800&display=swap');
:root{--bg:#0a0a0f;--surface:#111118;--surface2:#1a1a26;--border:#2a2a3a;
      --accent:#7c5cfc;--accent2:#00d4aa;--accent3:#ff6b6b;--text:#e8e8f0;--muted:#666680;}
html,body,[data-testid="stAppViewContainer"]{background:var(--bg)!important;color:var(--text)!important;font-family:'Syne',sans-serif!important;}
[data-testid="stSidebar"]{background:var(--surface)!important;border-right:1px solid var(--border)!important;}
[data-testid="stSidebar"] *{color:var(--text)!important;}
h1,h2,h3,h4{font-family:'Syne',sans-serif!important;}
.stButton>button{background:linear-gradient(135deg,var(--accent),#5a3fd4)!important;color:white!important;
  border:none!important;border-radius:8px!important;font-family:'Syne',sans-serif!important;
  font-weight:600!important;padding:.5rem 1.5rem!important;transition:all .2s!important;
  box-shadow:0 4px 15px rgba(124,92,252,.3)!important;}
.stButton>button:hover{transform:translateY(-1px)!important;box-shadow:0 6px 20px rgba(124,92,252,.5)!important;}
.metric-card{background:var(--surface2);border:1px solid var(--border);border-radius:12px;
  padding:1.2rem 1.5rem;text-align:center;position:relative;overflow:hidden;}
.metric-card::before{content:'';position:absolute;top:0;left:0;right:0;height:2px;
  background:linear-gradient(90deg,var(--accent),var(--accent2));}
.metric-value{font-size:2rem;font-weight:800;color:var(--accent2);font-family:'Space Mono',monospace;}
.metric-label{font-size:.75rem;color:var(--muted);text-transform:uppercase;letter-spacing:1px;margin-top:.3rem;}
.step-header{background:linear-gradient(135deg,var(--surface2),var(--surface));border:1px solid var(--border);
  border-left:3px solid var(--accent);border-radius:10px;padding:1rem 1.5rem;margin:1rem 0;}
.ai-card{background:linear-gradient(135deg,rgba(124,92,252,.1),rgba(0,212,170,.05));
  border:1px solid rgba(124,92,252,.3);border-radius:12px;padding:1.2rem;margin:.5rem 0;}
.feature-card{background:var(--surface2);border:1px solid var(--border);border-radius:12px;
  padding:1.2rem;text-align:center;transition:all .2s;}
.badge{display:inline-block;padding:.2rem .7rem;border-radius:20px;font-size:.7rem;font-weight:700;text-transform:uppercase;letter-spacing:.5px;}
.badge-purple{background:rgba(124,92,252,.2);color:#a78bfa;border:1px solid rgba(124,92,252,.3);}
.badge-green{background:rgba(0,212,170,.15);color:#00d4aa;border:1px solid rgba(0,212,170,.3);}
.badge-red{background:rgba(255,107,107,.15);color:#ff6b6b;border:1px solid rgba(255,107,107,.3);}
.badge-yellow{background:rgba(255,179,71,.15);color:#ffb347;border:1px solid rgba(255,179,71,.3);}
[data-testid="stDataFrame"]{border:1px solid var(--border)!important;border-radius:8px!important;}
.stProgress>div>div{background:linear-gradient(90deg,var(--accent),var(--accent2))!important;}
[data-testid="stExpander"]{background:var(--surface2)!important;border:1px solid var(--border)!important;border-radius:8px!important;}
.stTabs [data-baseweb="tab"]{background:transparent!important;color:var(--muted)!important;border-bottom:2px solid transparent!important;}
.stTabs [aria-selected="true"]{color:var(--accent)!important;border-bottom:2px solid var(--accent)!important;}
.footer-bar{background:var(--surface);border:1px solid var(--border);padding:1rem 2rem;margin-top:2.5rem;text-align:center;border-radius:12px;}
div[data-testid="stCheckbox"] label{color:var(--text)!important;}
.warn-card{background:linear-gradient(135deg,rgba(255,179,71,.1),rgba(255,107,107,.05));
  border:1px solid rgba(255,179,71,.4);border-radius:12px;padding:1rem 1.2rem;margin:.5rem 0;}
</style>
""", unsafe_allow_html=True)

#  Module imports 
from modules.file_loader import load_file
from modules.profiling import get_basic_info, get_column_summary, get_numeric_stats, get_categorical_stats
from modules.missing_handler import get_missing_summary, fill_missing_values, suggest_missing_strategy
from modules.duplicate_handler import remove_duplicates
from modules.exporter import export_data
from modules.clustering import (prepare_clustering_data, run_kmeans, run_dbscan,
                                run_agglomerative, reduce_to_2d, find_optimal_clusters,
                                get_best_clustering, run_all_clustering)
from modules.ai_recommender import recommend_clustering, generate_ai_report, recommend_model_export
from modules.model_export.export_manager import export_model

#  Session State ─
DEFAULTS = dict(df_raw=None, df_clean=None, df_cleaned_only=None, step=1,
                trained_models={}, results_df=None, best_model_name=None,
                problem_type=None, target_col=None, cluster_results=None,
                X_test=None, y_test=None, ai_report=None,
                missing_applied=False, dup_removed=False,
                preprocessing_log=[], training_dataset="preprocessed")
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

#  Helpers ─
PT = dict(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)',
          font=dict(color='#e8e8f0', family='Syne'),
          xaxis=dict(gridcolor='#2a2a3a', zerolinecolor='#2a2a3a'),
          yaxis=dict(gridcolor='#2a2a3a', zerolinecolor='#2a2a3a'),
          margin=dict(l=20, r=20, t=50, b=20))
COLORS = ["#7c5cfc","#00d4aa","#ff6b6b","#ffb347","#a78bfa","#34d399","#f87171","#fbbf24","#60a5fa","#e879f9"]

def mc(val, label):
    return f'<div class="metric-card"><div class="metric-value">{val}</div><div class="metric-label">{label}</div></div>'

def sh(icon, title, sub=""):
    st.markdown(f'''<div class="step-header"><span style="font-size:1.5rem">{icon}</span>
    <span style="font-size:1.2rem;font-weight:800;margin-left:.5rem">{title}</span>
    {"<br><span style='color:var(--muted);font-size:.85rem'>" + sub + "</span>" if sub else ""}
    </div>''', unsafe_allow_html=True)

def footer():
    st.markdown("""<div class="footer-bar">
    <span style="color:var(--muted);font-size:.8rem">
    ⚡ <strong style="color:var(--accent)">ML Studio Pro</strong> v2.0 &nbsp;|&nbsp;
    Upload → Clean → Train → Visualize → Export &nbsp;|&nbsp;
    <span style="color:var(--accent2)">Powered by Scikit-learn & Streamlit</span>
    </span></div>""", unsafe_allow_html=True)

def detect_problem_type(df, target_col):
    if target_col not in df.columns: return "classification"
    y = df[target_col].dropna()
    if y.dtype == "object": return "classification"
    if pd.api.types.is_numeric_dtype(y) and y.nunique() < max(10, int(0.05*len(y))):
        return "classification"
    return "regression"


# SIDEBAR

with st.sidebar:
    st.markdown("""<div style="text-align:center;padding:1rem 0 1.5rem">
    <div style="font-size:2.5rem">⚡</div>
    <div style="font-family:'Syne',sans-serif;font-size:1.3rem;font-weight:800;
        background:linear-gradient(135deg,#7c5cfc,#00d4aa);-webkit-background-clip:text;-webkit-text-fill-color:transparent">
        ML Studio Pro</div>
    <div style="font-size:.7rem;color:#666680;letter-spacing:2px;text-transform:uppercase">Advanced AutoML Platform</div>
    </div>""", unsafe_allow_html=True)

    STEPS = [("1","📂","Data Upload"),("2","🔍","Data Profiling"),
             ("3","🧹","Missing Values"),("4","♻️","Duplicates"),
             ("5","⚙️","Preprocessing"),("6","💾","Export Clean Data"),
             ("7","🤖","Model Training"),("8","📊","Cluster Visualizer"),
             ("9","📑","Report & Export")]
    st.markdown("### Pipeline Steps")
    for num, icon, label in STEPS:
        active = st.session_state.step == int(num)
        done   = st.session_state.step > int(num)
        col = "#7c5cfc" if active else ("#00d4aa" if done else "#2a2a3a")
        tc  = "#e8e8f0" if active else ("#00d4aa" if done else "#666680")
        bg  = "rgba(124,92,252,.1)" if active else "transparent"
        fw  = "700" if active else "400"
        st.markdown(f'''<div style="display:flex;align-items:center;gap:.7rem;padding:.5rem .7rem;
            border-radius:8px;background:{bg};border-left:3px solid {col};margin-bottom:.2rem">
            <span style="color:{col};font-size:.85rem">{icon}</span>
            <span style="color:{tc};font-size:.85rem;font-weight:{fw}">{label}</span>
        </div>''', unsafe_allow_html=True)

    if st.session_state.df_raw is not None:
        st.markdown("---")
        info = get_basic_info(st.session_state.df_raw)
        st.markdown("### Dataset")
        c1,c2 = st.columns(2)
        c1.metric("Rows", f"{info['rows']:,}")
        c2.metric("Cols", f"{info['columns']:,}")
        if st.session_state.df_clean is not None:
            st.caption(f"✅ Clean: {len(st.session_state.df_clean):,} rows")
    st.markdown("---")
    st.markdown("<div style='color:#666680;font-size:.7rem;text-align:center'>ML Studio Pro v2.0</div>", unsafe_allow_html=True)


# PAGE HEADER

st.markdown("""<div style="padding:1.5rem 0 .5rem">
<h1 style="font-family:'Syne',sans-serif;font-weight:800;font-size:2.2rem;margin:0;
    background:linear-gradient(135deg,#7c5cfc,#00d4aa);-webkit-background-clip:text;-webkit-text-fill-color:transparent">
    ⚡ ML Studio Pro</h1>
<p style="color:#666680;margin:.3rem 0 0">End-to-end ML Pipeline — Upload → Clean → Preprocess → Train → Visualize → Export</p>
</div>""", unsafe_allow_html=True)


# STEP 1 — UPLOAD

if st.session_state.step == 1:
    sh("📂", "Upload Your Dataset", "CSV, Excel, JSON, XML, YAML, SQLite — all supported")
    col_up, col_demo = st.columns([3, 1])
    with col_up:
        uploaded = st.file_uploader("Drop your file here",
            type=["csv","xlsx","xls","json","xml","yaml","yml","db"])
    with col_demo:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🎲 Load Demo Data"):
            from sklearn.datasets import load_iris
            iris = load_iris(as_frame=True); df_d = iris.frame
            st.session_state.df_raw = df_d; st.session_state.df_clean = df_d.copy()
            st.session_state.step = 2; st.rerun()

    if uploaded:
        with st.spinner("Reading file..."):
            df, msg = load_file(uploaded)
        if df is not None:
            st.session_state.df_raw = df; st.session_state.df_clean = df.copy()
            st.success(f"✅ Loaded **{uploaded.name}** — {len(df):,} rows × {len(df.columns)} columns")
            st.markdown("#### 👁️ Data Preview")
            st.dataframe(df.head(10), use_container_width=True)
            info = get_basic_info(df)
            c1,c2,c3,c4 = st.columns(4)
            c1.markdown(mc(f"{info['rows']:,}", "Total Rows"), unsafe_allow_html=True)
            c2.markdown(mc(f"{info['columns']:,}", "Columns"), unsafe_allow_html=True)
            c3.markdown(mc(f"{int(df.isna().sum().sum()):,}", "Missing"), unsafe_allow_html=True)
            c4.markdown(mc(f"{info['duplicate_rows']:,}", "Duplicates"), unsafe_allow_html=True)
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("Next → Explore Data ▶"): st.session_state.step = 2; st.rerun()
        else:
            st.error(f"❌ {msg}")

    st.markdown("---")
    st.markdown("""<div style="text-align:center;margin-bottom:1.5rem">
    <span style="font-size:1.4rem;font-weight:800;background:linear-gradient(135deg,#7c5cfc,#00d4aa);
    -webkit-background-clip:text;-webkit-text-fill-color:transparent">🚀 How ML Studio Pro Works</span><br>
    <span style="color:var(--muted);font-size:.9rem">Complete AutoML pipeline — no coding required</span>
    </div>""", unsafe_allow_html=True)

    features = [
        ("📂","Upload Any File","CSV, Excel, JSON, XML, YAML, SQLite"),
        ("🔍","Smart Profiling","Distributions, correlations & missing patterns"),
        ("🧹","AI-Guided Cleaning","Per-column fill strategies suggested by AI"),
        ("♻️","Duplicate Detection","Full-row & key-column safe removal"),
        ("⚙️","Preprocessing","Encode, scale, normalize, drop columns"),
        ("💾","Multi-format Export","CSV, Excel, JSON, XML, YAML, SQLite"),
        ("🤖","AutoML Training","6 classifiers & 4 regressors benchmarked"),
        ("📊","Cluster Visualizer","5 chart types with full axis control"),
        ("📑","Pro Reports","Excel, PDF & Word with charts & AI insights"),
    ]
    cols = st.columns(3)
    for i,(icon,title,desc) in enumerate(features):
        with cols[i%3]:
            st.markdown(f"""<div class="feature-card" style="margin-bottom:1rem">
            <div style="font-size:1.8rem;margin-bottom:.5rem">{icon}</div>
            <div style="font-weight:700;color:var(--accent2);font-size:.9rem;margin-bottom:.4rem">{title}</div>
            <div style="font-size:.75rem;color:var(--muted);line-height:1.5">{desc}</div>
            </div>""", unsafe_allow_html=True)

    st.markdown("""<div class="ai-card">
    <strong>🤖 AI-Powered Throughout:</strong> Every step has built-in AI analysis — fill strategies,
    model selection, clustering recommendations, and full session reports. All explained clearly.
    </div>""", unsafe_allow_html=True)
    footer()


# STEP 2 — PROFILING

elif st.session_state.step == 2:
    df = st.session_state.df_clean
    sh("🔍", "Data Profiling", "Understand your dataset before cleaning")
    info = get_basic_info(df)
    c1,c2,c3,c4 = st.columns(4)
    c1.markdown(mc(f"{info['rows']:,}", "Rows"), unsafe_allow_html=True)
    c2.markdown(mc(f"{info['columns']:,}", "Columns"), unsafe_allow_html=True)
    c3.markdown(mc(f"{int(df.isna().sum().sum()):,}", "Missing"), unsafe_allow_html=True)
    c4.markdown(mc(f"{info['duplicate_rows']:,}", "Duplicates"), unsafe_allow_html=True)
    st.markdown("---")
    t1,t2,t3,t4 = st.tabs(["📋 Column Summary","📈 Numeric Stats","🔤 Categorical Stats","🗺️ Missing Chart"])
    with t1:
        st.dataframe(get_column_summary(df), use_container_width=True, height=400)
    with t2:
        ns = get_numeric_stats(df)
        if not ns.empty:
            st.dataframe(ns, use_container_width=True, height=350)
            num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
            if num_cols:
                sel = st.selectbox("Column for distribution", num_cols, key="dist_col")
                fig = px.histogram(df, x=sel, nbins=40, color_discrete_sequence=["#7c5cfc"])
                fig.update_layout(**PT, title=f"Distribution — {sel}")
                st.plotly_chart(fig, use_container_width=True)
        else: st.info("No numeric columns.")
    with t3:
        cs = get_categorical_stats(df)
        if not cs.empty: st.dataframe(cs, use_container_width=True, height=300)
        else: st.info("No categorical columns.")
    with t4:
        mp = (df.isna().sum()/len(df)*100).reset_index()
        mp.columns = ["Column","Missing %"]; mp = mp[mp["Missing %"]>0]
        if not mp.empty:
            fig = px.bar(mp, x="Missing %", y="Column", orientation="h",
                         color="Missing %", color_continuous_scale=["#7c5cfc","#ff6b6b"])
            fig.update_layout(**PT, title="Missing Values by Column")
            st.plotly_chart(fig, use_container_width=True)
        else: st.success("🎉 No missing values!")
    num_df = df.select_dtypes(include=[np.number])
    if len(num_df.columns) > 1:
        st.markdown("#### 🔗 Correlation Heatmap")
        corr = num_df.corr()
        fig = go.Figure(go.Heatmap(z=corr.values, x=corr.columns, y=corr.index,
            colorscale=[[0,"#ff6b6b"],[.5,"#1a1a26"],[1,"#7c5cfc"]],
            text=np.round(corr.values,2), texttemplate="%{text}"))
        fig.update_layout(**PT, title="Correlation Matrix"); st.plotly_chart(fig, use_container_width=True)
    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=1; st.rerun()
    with c2:
        if st.button("Next → Missing Values ▶"): st.session_state.step=3; st.rerun()
    footer()


# STEP 3 — MISSING VALUES

elif st.session_state.step == 3:
    df = st.session_state.df_clean
    sh("🧹", "Handle Missing Values", "Choose strategy per column")
    miss_df = get_missing_summary(df)
    if miss_df.empty:
        st.markdown("""<div class="ai-card"><span style="font-size:1.5rem">🎉</span>
        <strong> No missing values!</strong> Dataset is already clean.</div>""", unsafe_allow_html=True)
    else:
        st.markdown(f"**{len(miss_df)} column(s) have missing values:**")
        st.dataframe(miss_df, use_container_width=True)
        suggested = suggest_missing_strategy(df)
        st.markdown("#### ⚙️ Configure Fill Strategy")
        st.markdown("""<div class="ai-card">🤖 <strong>AI suggested strategies</strong>
        based on data skewness & type. Override any as needed.</div>""", unsafe_allow_html=True)
        strategies, custom_vals = {}, {}
        for _, row in miss_df.iterrows():
            col = row["column"]; dtype = row["dtype"]
            is_num = ("int" in dtype or "float" in dtype)
            sug = suggested.get(col, "mean" if is_num else "unknown")
            opts = (["mean","median","mode","ffill","bfill","zero","constant"]
                    if is_num else ["unknown","mode","ffill","bfill","empty","constant"])
            idx = opts.index(sug) if sug in opts else 0
            with st.expander(f"📌 `{col}` — {row['missing_%']}% missing ({dtype})", expanded=True):
                cc1,cc2 = st.columns([2,3])
                with cc1:
                    badge = "badge-red" if row["missing_%"]>30 else "badge-yellow"
                    st.markdown(f'<span class="badge {badge}">{row["missing_%"]}% missing</span>'
                                f'<span class="badge badge-purple" style="margin-left:.3rem">{dtype}</span>',
                                unsafe_allow_html=True)
                    choice = st.selectbox("Strategy", opts, index=idx, key=f"s_{col}")
                    strategies[col] = choice
                    if choice == "constant":
                        cv = st.text_input(f"Value for `{col}`", key=f"cv_{col}")
                        if cv:
                            try: custom_vals[col] = float(cv) if is_num else cv
                            except: custom_vals[col] = cv
                with cc2:
                    sample = df[col].dropna().head(5).tolist()
                    st.markdown(f"**Sample:** `{sample}`")
                    st.markdown(f"🤖 AI suggestion: **`{sug}`**")
                    st.caption(f"Missing: {int(row['missing_count'])} of {len(df)} rows")
        if st.button("✅ Apply Strategies & Fill"):
            try:
                df_filled = fill_missing_values(df, strategies, custom_vals)
                remaining = int(df_filled.isna().sum().sum())
                st.session_state.df_clean = df_filled
                st.session_state.missing_applied = True
                st.session_state.preprocessing_log.append("Missing values filled")
                if remaining == 0: st.success("✅ All missing values filled!")
                else: st.warning(f"⚠️ {remaining} values remain — check constant fields.")
                st.rerun()
            except Exception as e: st.error(f"❌ {e}")
    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=2; st.rerun()
    with c2:
        if st.button("Next → Duplicates ▶"): st.session_state.step=4; st.rerun()
    footer()


# STEP 4 — DUPLICATES

elif st.session_state.step == 4:
    df = st.session_state.df_clean
    sh("♻️", "Handle Duplicates", "Detect and remove duplicate rows")
    exact_dup = int(df.duplicated(keep="first").sum())
    pct = round(exact_dup/len(df)*100, 2) if len(df)>0 else 0
    c1,c2,c3 = st.columns(3)
    c1.markdown(mc(f"{exact_dup:,}", "Removable Duplicates"), unsafe_allow_html=True)
    c2.markdown(mc(f"{pct}%", "Duplicate %"), unsafe_allow_html=True)
    c3.markdown(mc(f"{len(df)-exact_dup:,}", "Rows After Clean"), unsafe_allow_html=True)
    if exact_dup == 0:
        st.markdown("""<div class="ai-card"><span style="font-size:1.5rem">✅</span>
        <strong> No duplicate rows found!</strong> Every row is unique.</div>""", unsafe_allow_html=True)
    else:
        with st.expander("👀 Preview Duplicate Rows", expanded=False):
            st.dataframe(df[df.duplicated(keep=False)].head(50), use_container_width=True)
        cc1,cc2 = st.columns(2)
        with cc1:
            keep_opt = st.selectbox("Which copy to keep?",
                ["first — keep first occurrence","last  — keep last occurrence","none  — remove ALL copies"])
            keep_val = (False if "none" in keep_opt else "first" if "first" in keep_opt else "last")
        with cc2:
            use_subset = st.checkbox("Check only specific columns", value=False)
            subset = None
            if use_subset:
                sel_cols = st.multiselect("Columns to check", df.columns.tolist())
                subset = sel_cols if sel_cols else None
                if subset:
                    sub_dup = int(df.duplicated(subset=subset, keep="first").sum())
                    st.info(f"With selected columns: **{sub_dup}** duplicates")
        st.markdown("""<div class="ai-card">🤖 <strong>AI Tip:</strong>
        Full-row comparison is safest. Use column subset only for key-based deduplication.</div>""",
        unsafe_allow_html=True)
        if st.button("🗑️ Remove Duplicate Rows"):
            before = len(df)
            df_clean = df.drop_duplicates(subset=subset, keep=keep_val)
            after = len(df_clean)
            st.session_state.df_clean = df_clean
            st.session_state.dup_removed = True
            st.session_state.preprocessing_log.append(f"Removed {before-after} duplicate rows")
            st.success(f"✅ Removed **{before-after}** rows. {before:,} → {after:,}")
            time.sleep(0.3); st.rerun()
    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=3; st.rerun()
    with c2:
        if st.button("Next → Preprocessing ▶"):
            # Snapshot the cleaned-only df before any preprocessing
            st.session_state.df_cleaned_only = st.session_state.df_clean.copy()
            st.session_state.step=5; st.rerun()
    footer()


# STEP 5 — PREPROCESSING (NEW)

elif st.session_state.step == 5:
    df = st.session_state.df_clean
    sh("⚙️", "Preprocessing", "Encode, scale, normalize & engineer your features")

    #  Smart AI Guide ─
    st.markdown("### 🤖 AI Preprocessing Recommendations")
    ai_rows = []
    for col in df.columns:
        dtype = str(df[col].dtype)
        n_unique = df[col].nunique()
        n_total  = len(df)
        is_num   = pd.api.types.is_numeric_dtype(df[col])
        is_cat   = dtype in ("object","category") or not is_num

        if is_num:
            col_data = df[col].dropna()
            try:
                rng   = float(col_data.max() - col_data.min())
                skew  = float(col_data.skew())
                std   = float(col_data.std())
            except Exception:
                rng = skew = std = 0

            if rng > 100 or std > 10:
                if abs(skew) > 1:
                    action = "RobustScaler"
                    reason = f"High skew ({skew:.2f}) + large range ({rng:.0f}) — robust to outliers"
                else:
                    action = "StandardScaler"
                    reason = f"Large range ({rng:.0f}), std={std:.2f} — normalize to zero mean"
            elif rng > 1:
                action = "MinMaxScaler"
                reason = f"Moderate range ({rng:.2f}) — scale to 0-1"
            else:
                action = "Optional / Skip"
                reason = "Already small range — scaling may not be needed"
        else:
            if n_unique > 50:
                action = "Drop or Hash"
                reason = f"Very high cardinality ({n_unique} unique) — likely an ID or free-text, drop it"
            elif n_unique > 10:
                action = "Label Encoding"
                reason = f"High cardinality ({n_unique} unique) — label encoding preferred"
            elif n_unique > 2:
                action = "One-Hot Encoding"
                reason = f"Low cardinality ({n_unique} unique) — safe for one-hot"
            else:
                action = "Label Encoding"
                reason = f"Binary / bool ({n_unique} unique) — simple label encoding"

        ai_rows.append({"Column": col, "Type": dtype, "Unique Values": n_unique,
                        "Recommended Action": action, "Reason": reason})

    ai_guide_df = pd.DataFrame(ai_rows)

    def _style_action(val):
        if "Drop" in val:   return "background:rgba(255,107,107,.15);color:#ff6b6b"
        if "Scaler" in val: return "background:rgba(0,212,170,.15);color:#00d4aa"
        if "Encoding" in val: return "background:rgba(124,92,252,.15);color:#a78bfa"
        return "background:rgba(255,179,71,.1);color:#ffb347"

    styled = ai_guide_df.style.applymap(_style_action, subset=["Recommended Action"])
    st.dataframe(styled, use_container_width=True, height=min(35*len(ai_guide_df)+40, 380))
    st.markdown("""<div class="ai-card" style="font-size:.82rem">
    ⚡ <strong>Preprocessing is optional.</strong> The AI recommendations above are suggestions based on your data.
    Use the tabs below to apply any steps you want — then click <strong>Next</strong> to continue.
    The original cleaned data is always saved separately for download.
    </div>""", unsafe_allow_html=True)
    st.markdown("---")

    all_cols   = df.columns.tolist()
    num_cols   = df.select_dtypes(include=[np.number]).columns.tolist()
    cat_cols   = df.select_dtypes(include=["object","category"]).columns.tolist()

    tab1, tab2, tab3, tab4 = st.tabs(["🔢 Encoding","📏 Scaling / Normalization","🗑️ Drop Columns","🔬 Feature Info"])

    with tab1:
        st.markdown("#### Encode Categorical Columns")
        if not cat_cols:
            st.info("No categorical columns found.")
        else:
            enc_method = st.selectbox("Encoding method",
                ["Label Encoding","One-Hot Encoding","Ordinal Encoding"], key="enc_method")
            enc_cols = st.multiselect("Select columns to encode", cat_cols, default=cat_cols, key="enc_cols")
            if st.button("✅ Apply Encoding"):
                try:
                    df_enc = st.session_state.df_clean.copy()
                    log_msg = []
                    if enc_method == "Label Encoding":
                        from sklearn.preprocessing import LabelEncoder
                        le = LabelEncoder()
                        for col in enc_cols:
                            if col in df_enc.columns:
                                df_enc[col] = le.fit_transform(df_enc[col].astype(str))
                        log_msg.append(f"Label encoded: {enc_cols}")
                    elif enc_method == "One-Hot Encoding":
                        df_enc = pd.get_dummies(df_enc, columns=enc_cols, drop_first=False)
                        log_msg.append(f"One-hot encoded: {enc_cols}")
                    elif enc_method == "Ordinal Encoding":
                        from sklearn.preprocessing import OrdinalEncoder
                        oe = OrdinalEncoder()
                        df_enc[enc_cols] = oe.fit_transform(df_enc[enc_cols].astype(str))
                        log_msg.append(f"Ordinal encoded: {enc_cols}")
                    st.session_state.df_clean = df_enc
                    st.session_state.preprocessing_log.extend(log_msg)
                    st.success(f"✅ {enc_method} applied to {len(enc_cols)} columns!")
                    st.dataframe(df_enc.head(5), use_container_width=True)
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Encoding error: {e}")

    with tab2:
        st.markdown("#### Scale / Normalize Numeric Columns")
        if not num_cols:
            st.info("No numeric columns found.")
        else:
            scale_method = st.selectbox("Scaling method",
                ["StandardScaler (mean=0, std=1)",
                 "MinMaxScaler (0 to 1)",
                 "RobustScaler (outlier-resistant)",
                 "Normalizer (unit norm per row)"], key="scale_method")
            scale_cols = st.multiselect("Columns to scale", num_cols, default=num_cols[:3] if len(num_cols)>3 else num_cols, key="scale_cols")
            if st.button("✅ Apply Scaling"):
                try:
                    from sklearn.preprocessing import StandardScaler, MinMaxScaler, RobustScaler, Normalizer
                    df_sc = st.session_state.df_clean.copy()
                    if scale_cols:
                        method_map = {
                            "StandardScaler (mean=0, std=1)": StandardScaler(),
                            "MinMaxScaler (0 to 1)": MinMaxScaler(),
                            "RobustScaler (outlier-resistant)": RobustScaler(),
                            "Normalizer (unit norm per row)": Normalizer(),
                        }
                        scaler = method_map[scale_method]
                        df_sc[scale_cols] = scaler.fit_transform(df_sc[scale_cols].fillna(0))
                        st.session_state.df_clean = df_sc
                        st.session_state.preprocessing_log.append(f"{scale_method} on {scale_cols}")
                        st.success(f"✅ Scaling applied to {len(scale_cols)} columns!")
                        st.dataframe(df_sc[scale_cols].head(5), use_container_width=True)
                except Exception as e:
                    st.error(f"❌ Scaling error: {e}")

    with tab3:
        st.markdown("#### Drop Columns")
        drop_cols = st.multiselect("Select columns to drop", st.session_state.df_clean.columns.tolist(), key="drop_cols")
        if drop_cols:
            st.warning(f"⚠️ Will drop: **{drop_cols}**")
            if st.button("🗑️ Drop Selected Columns"):
                df_drop = st.session_state.df_clean.drop(columns=drop_cols)
                st.session_state.df_clean = df_drop
                st.session_state.preprocessing_log.append(f"Dropped columns: {drop_cols}")
                st.success(f"✅ Dropped {len(drop_cols)} column(s). Remaining: {len(df_drop.columns)}")
                st.rerun()

    with tab4:
        st.markdown("#### 📊 Current Dataset Info After Preprocessing")
        df_now = st.session_state.df_clean
        info = get_basic_info(df_now)
        c1,c2,c3 = st.columns(3)
        c1.markdown(mc(f"{info['rows']:,}", "Rows"), unsafe_allow_html=True)
        c2.markdown(mc(f"{info['columns']:,}", "Columns"), unsafe_allow_html=True)
        c3.markdown(mc(f"{int(df_now.isna().sum().sum()):,}", "Missing"), unsafe_allow_html=True)
        st.dataframe(get_column_summary(df_now), use_container_width=True, height=300)
        if st.session_state.preprocessing_log:
            st.markdown("#### 📋 Preprocessing Log")
            for i,log in enumerate(st.session_state.preprocessing_log):
                st.markdown(f"`{i+1}.` {log}")

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=4; st.rerun()
    with c2:
        if st.button("Next → Export Clean Data ▶"): st.session_state.step=6; st.rerun()
    footer()


# STEP 6 — EXPORT CLEAN DATA (dual download + training selector)

elif st.session_state.step == 6:
    sh("💾", "Export Clean Dataset", "Download cleaned & preprocessed data — choose which to use for training")

    df_preprocessed = st.session_state.df_clean
    df_cleaned      = st.session_state.df_cleaned_only if st.session_state.df_cleaned_only is not None else df_preprocessed

    EXT  = {"csv":"csv","excel":"xlsx","json":"json","xml":"xml","yaml":"yaml","sql":"db"}
    MIME = {"csv":"text/csv","excel":"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "json":"application/json","xml":"application/xml","yaml":"text/yaml","sql":"application/octet-stream"}
    FMT_LABELS = {"csv":"📄 CSV","excel":"📊 Excel (.xlsx)","json":"📋 JSON",
                  "xml":"🏷️ XML","yaml":"⚙️ YAML","sql":"🗄️ SQLite DB"}

    def _make_download_bytes(df, fmt):
        try:
            if fmt == "sql":
                import sqlite3, tempfile
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".db"); tmp.close()
                sdf = df.copy()
                sdf.columns = [str(c).replace(" ","_").replace(".","_").replace("-","_") for c in sdf.columns]
                conn = sqlite3.connect(tmp.name)
                sdf.to_sql("data", conn, if_exists="replace", index=False)
                conn.commit(); conn.close()
                with open(tmp.name,"rb") as f: raw = f.read()
                os.remove(tmp.name)
                return raw
            else:
                raw, _ = export_data(df, fmt)
                return raw.encode("utf-8") if isinstance(raw, str) else raw
        except Exception as ex:
            st.error(f"❌ Export error: {ex}"); return None

    col_a, col_b = st.columns(2)

    #  LEFT: Cleaned-Only ─
    with col_a:
        st.markdown("""<div class="metric-card" style="margin-bottom:1rem">
        <div style="font-size:1.8rem">🧹</div>
        <div style="font-weight:700;color:var(--accent2);font-size:1rem">Cleaned Data</div>
        <div style="font-size:.75rem;color:var(--muted);margin-top:.3rem">After missing/duplicate fix — before encoding & scaling</div>
        </div>""", unsafe_allow_html=True)
        st.markdown(f"`{len(df_cleaned):,}` rows × `{len(df_cleaned.columns)}` columns")
        st.dataframe(df_cleaned.head(5), use_container_width=True)
        fmt_a = st.selectbox("Format", list(EXT.keys()), format_func=lambda x: FMT_LABELS.get(x,x), key="fmt_cleaned")
        if st.button("⬇️ Download Cleaned Data", use_container_width=True, key="dl_cleaned"):
            with st.spinner("Preparing..."):
                raw = _make_download_bytes(df_cleaned, fmt_a)
                if raw:
                    st.download_button(f"💾 cleaned_data.{EXT[fmt_a]}", data=raw,
                        file_name=f"cleaned_data.{EXT[fmt_a]}", mime=MIME[fmt_a], key="dlb_cleaned")
                    st.success("✅ Ready!")

    #  RIGHT: Preprocessed ─
    with col_b:
        st.markdown("""<div class="metric-card" style="margin-bottom:1rem">
        <div style="font-size:1.8rem">⚙️</div>
        <div style="font-weight:700;color:var(--accent);font-size:1rem">Preprocessed Data</div>
        <div style="font-size:.75rem;color:var(--muted);margin-top:.3rem">After encoding, scaling & all preprocessing steps</div>
        </div>""", unsafe_allow_html=True)
        st.markdown(f"`{len(df_preprocessed):,}` rows × `{len(df_preprocessed.columns)}` columns")
        st.dataframe(df_preprocessed.head(5), use_container_width=True)
        fmt_b = st.selectbox("Format", list(EXT.keys()), format_func=lambda x: FMT_LABELS.get(x,x), key="fmt_preprocessed")
        if st.button("⬇️ Download Preprocessed Data", use_container_width=True, key="dl_preprocessed"):
            with st.spinner("Preparing..."):
                raw = _make_download_bytes(df_preprocessed, fmt_b)
                if raw:
                    st.download_button(f"💾 preprocessed_data.{EXT[fmt_b]}", data=raw,
                        file_name=f"preprocessed_data.{EXT[fmt_b]}", mime=MIME[fmt_b], key="dlb_preprocessed")
                    st.success("✅ Ready!")

    #  Training Dataset Selector ─
    st.markdown("---")
    st.markdown("### 🎯 Dataset to Use for Model Training")
    st.markdown("""<div class="ai-card">
    🤖 <strong>AI Tip:</strong> Use <strong>Preprocessed Data</strong> if you applied encoding/scaling —
    models require numeric inputs. Use <strong>Cleaned Data</strong> if you prefer the raw cleaned format
    (the training step will auto-encode categoricals).
    </div>""", unsafe_allow_html=True)
    train_choice = st.radio(
        "Which dataset should Step 7 (Model Training) use?",
        ["preprocessed", "cleaned"],
        index=0 if st.session_state.training_dataset == "preprocessed" else 1,
        format_func=lambda x: "⚙️ Preprocessed Data (recommended)" if x=="preprocessed" else "🧹 Cleaned Data (raw)",
        horizontal=True, key="train_ds_radio"
    )
    st.session_state.training_dataset = train_choice

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=5; st.rerun()
    with c2:
        if st.button("Next → Train Models ▶"): st.session_state.step=7; st.rerun()
    footer()


# STEP 7 — MODEL TRAINING + RESULTS

elif st.session_state.step == 7:
    # Use dataset chosen in Step 6
    _use_prep = st.session_state.training_dataset == "preprocessed"
    df = st.session_state.df_clean if _use_prep else (
        st.session_state.df_cleaned_only if st.session_state.df_cleaned_only is not None
        else st.session_state.df_clean)
    sh("🤖", "Model Training & Results", "Train, evaluate and download your ML model")

    # Dataset badge
    _ds_label = "⚙️ Preprocessed" if _use_prep else "🧹 Cleaned"
    st.markdown(f"""<div class="ai-card" style="font-size:.82rem;padding:.7rem 1rem">
    🎯 <strong>Training on:</strong> <span class="badge badge-green">{_ds_label}</span>
    &nbsp;— {len(df):,} rows × {len(df.columns)} columns
    &nbsp;<span style="color:var(--muted);font-size:.75rem">(change in Step 6)</span>
    </div>""", unsafe_allow_html=True)

    all_cols = df.columns.tolist()

    #  AI Analysis FIRST ─
    # Use last target col if set, else default to last column
    _default_target = st.session_state.target_col if st.session_state.target_col in all_cols else all_cols[-1]
    _ai_target = _default_target  # temp for AI analysis

    if st.session_state.ai_report is None or st.session_state.target_col != _ai_target:
        with st.spinner("🤖 AI analysing dataset..."):
            report = generate_ai_report(df, _ai_target)
        st.session_state.ai_report = report
        st.session_state.target_col = _ai_target
    report = st.session_state.ai_report

    with st.expander("🤖 AI Analysis & Recommendations", expanded=True):
        cc = st.columns(4)
        cc[0].markdown(mc(f"{report['dataset']['rows']:,}", "Rows"), unsafe_allow_html=True)
        cc[1].markdown(mc(f"{report['dataset']['columns']:,}", "Columns"), unsafe_allow_html=True)
        cc[2].markdown(mc(f"{report['dataset']['missing_cells']:,}", "Missing"), unsafe_allow_html=True)
        cc[3].markdown(mc(f"{report['dataset']['duplicate_rows']:,}", "Duplicates"), unsafe_allow_html=True)
        pt_ai      = report.get("problem_type") or detect_problem_type(df, _ai_target)
        rec_models = report.get("recommended_models", [])
        imp_feat   = report.get("important_features", [])
        clust_rec  = report.get("clustering", {})
        miss_strat = report.get("missing_strategy", {})
        ptc = "badge-purple" if pt_ai=="classification" else "badge-green"
        st.markdown(f"""<div class="ai-card" style="margin-top:1rem">
        <strong>🤖 AI Complete Recommendation</strong><br><br>
        📌 <strong>Problem Type:</strong> <span class="badge {ptc}">{pt_ai or "N/A"}</span><br><br>
        🧠 <strong>Recommended Models:</strong> {", ".join(rec_models[:5]) or "N/A"}<br><br>
        ⭐ <strong>Top Features (by correlation):</strong> {", ".join([str(f) for f in imp_feat[:5]]) or "N/A"}<br><br>
        🔵 <strong>Best Clustering:</strong> {clust_rec.get("recommended","N/A")} — {clust_rec.get("reason","")}<br><br>
        🧹 <strong>Missing Strategy:</strong> {", ".join([f"{k}→{v}" for k,v in list(miss_strat.items())[:4]]) or "No missing data"}
        </div>""", unsafe_allow_html=True)

    st.markdown("---")

    #  Target Column + Task Type (AFTER AI) ─
    target_col = st.selectbox("🎯 Target Column (what to predict)", all_cols,
        index=all_cols.index(_default_target) if _default_target in all_cols else len(all_cols)-1)

    # Refresh AI if target changed
    if target_col != st.session_state.target_col:
        with st.spinner("🤖 AI re-analysing..."):
            report = generate_ai_report(df, target_col)
        st.session_state.ai_report = report
        st.session_state.target_col = target_col
        report = st.session_state.ai_report

    auto_type = detect_problem_type(df, target_col)

    #  FIXED: Only classification & regression — no clustering here 
    task_opts = ["classification", "regression"]
    task_type = st.selectbox("📌 Task Type", task_opts,
        index=task_opts.index(auto_type) if auto_type in task_opts else 0)

    # Show warning if AI says regression but user picks classification or vice versa
    if auto_type != task_type:
        st.markdown(f"""<div class="warn-card">
        ⚠️ <strong>Note:</strong> AI detected <span class="badge badge-green">{auto_type}</span>
        for this target column, but you selected <span class="badge badge-purple">{task_type}</span>.
        This may produce unexpected results — make sure your target column matches the task type.
        </div>""", unsafe_allow_html=True)
    else:
        st.markdown(f"🤖 AI detected: <span class='badge badge-purple'>{auto_type}</span>", unsafe_allow_html=True)

    st.markdown("#### 🎛️ Select Models to Train")
    if task_type == "classification":
        MODEL_KEYS = ["LogisticRegression","DecisionTree","RandomForest","GradientBoosting","SVC","KNN","NaiveBayes"]
        MODEL_DISP = ["Logistic Regression","Decision Tree","Random Forest","Gradient Boosting","SVC","KNN","Naive Bayes"]
        st.markdown("""<div class="ai-card" style="font-size:.85rem">
        📋 <strong>Classification models:</strong> Predict categories/classes. <strong>Gradient Boosting</strong> usually achieves highest accuracy.</div>""",
        unsafe_allow_html=True)
    else:
        MODEL_KEYS = ["LinearRegression","DecisionTree","RandomForest","GradientBoosting","ExtraTrees","SVR"]
        MODEL_DISP = ["Linear Regression","Decision Tree","Random Forest","Gradient Boosting","Extra Trees","SVR"]
        st.markdown("""<div class="ai-card" style="font-size:.85rem">
        📋 <strong>Regression models:</strong> <strong>Gradient Boosting</strong> and <strong>Extra Trees</strong> typically achieve R²≥0.90 on structured data.</div>""",
        unsafe_allow_html=True)

    sel_keys = []
    cols3 = st.columns(3)
    for i,(disp,key) in enumerate(zip(MODEL_DISP,MODEL_KEYS)):
        with cols3[i%3]:
            if st.checkbox(disp, value=True, key=f"m_{key}_{task_type}"):
                sel_keys.append(key)

    if st.button("🚀 Train Models", use_container_width=True):
        if not sel_keys: st.error("Select at least one model."); st.stop()
        try:
            from sklearn.model_selection import train_test_split, StratifiedShuffleSplit
            from sklearn.preprocessing import StandardScaler
            from sklearn.pipeline import Pipeline
            from sklearn.linear_model import LogisticRegression, LinearRegression
            from sklearn.tree import DecisionTreeClassifier, DecisionTreeRegressor
            from sklearn.ensemble import (RandomForestClassifier, RandomForestRegressor,
                                          GradientBoostingClassifier, GradientBoostingRegressor,
                                          ExtraTreesClassifier, ExtraTreesRegressor)
            from sklearn.svm import SVC, SVR
            from sklearn.neighbors import KNeighborsClassifier
            from sklearn.naive_bayes import GaussianNB
            from sklearn.metrics import (accuracy_score, precision_score, recall_score,
                                        f1_score, r2_score, mean_absolute_error, mean_squared_error)

            df_prep = df.copy().dropna()
            X = pd.get_dummies(df_prep.drop(columns=[target_col]), drop_first=True)
            y = df_prep[target_col]

            if X.empty:
                st.error("❌ No feature columns left after dropping target. Please check your data."); st.stop()

            stratify = None
            if task_type == "classification":
                counts  = y.value_counts()
                bad_cls = counts[counts < 2].index.tolist()
                if bad_cls:
                    st.warning(f"⚠️ Classes with <2 samples removed: {bad_cls}")
                    mask = ~y.isin(bad_cls); X,y = X[mask], y[mask]
                n_samples = len(y); n_classes = y.nunique()
                # Dynamic test_size
                min_test  = max(n_classes, 2)
                test_frac = max(0.1, min(0.3, min_test / n_samples))
                if int(n_samples * test_frac) < n_classes:
                    test_frac = min(0.5, (n_classes * 1.5) / n_samples)
                # Feasibility check for stratify
                try:
                    sss = StratifiedShuffleSplit(n_splits=1, test_size=test_frac, random_state=42)
                    list(sss.split(X, y)); stratify = y
                except Exception: stratify = None
            else:
                n_samples = len(y); test_frac = 0.2

            if int(n_samples * test_frac) < 1 or n_samples < 4:
                st.error(f"❌ Dataset too small ({n_samples} rows) to train. Need ≥4 rows."); st.stop()

            X_train,X_test,y_train,y_test = train_test_split(
                X, y, test_size=test_frac, random_state=42, stratify=stratify)

            ALL_CLF = {
                "LogisticRegression": Pipeline([("sc",StandardScaler()),("m",LogisticRegression(max_iter=2000))]),
                "DecisionTree":       DecisionTreeClassifier(random_state=42),
                "RandomForest":       RandomForestClassifier(n_estimators=200, random_state=42),
                "GradientBoosting":   GradientBoostingClassifier(n_estimators=200, learning_rate=0.05, random_state=42),
                "SVC":                Pipeline([("sc",StandardScaler()),("m",SVC(probability=True))]),
                "KNN":                Pipeline([("sc",StandardScaler()),("m",KNeighborsClassifier())]),
                "NaiveBayes":         GaussianNB(),
            }
            ALL_REG = {
                "LinearRegression": Pipeline([("sc",StandardScaler()),("m",LinearRegression())]),
                "DecisionTree":     DecisionTreeRegressor(random_state=42),
                "RandomForest":     RandomForestRegressor(n_estimators=200, random_state=42),
                "GradientBoosting": GradientBoostingRegressor(n_estimators=200, learning_rate=0.05, random_state=42),
                "ExtraTrees":       ExtraTreesRegressor(n_estimators=200, random_state=42),
                "SVR":              Pipeline([("sc",StandardScaler()),("m",SVR(kernel="rbf"))]),
            }
            pool = {k:v for k,v in (ALL_CLF if task_type=="classification" else ALL_REG).items()
                    if k in sel_keys}

            results, trained_ms = [], {}
            prog = st.progress(0, text="Training models...")
            for i,(name,model) in enumerate(pool.items()):
                try:
                    model.fit(X_train, y_train); y_pred = model.predict(X_test)
                    trained_ms[name] = model
                    if task_type=="classification":
                        results.append({"Model":name,
                            "Accuracy": round(accuracy_score(y_test,y_pred),4),
                            "Precision":round(precision_score(y_test,y_pred,average="weighted",zero_division=0),4),
                            "Recall":   round(recall_score(y_test,y_pred,average="weighted",zero_division=0),4),
                            "F1 Score": round(f1_score(y_test,y_pred,average="weighted",zero_division=0),4)})
                    else:
                        r2 = max(-1.0, round(r2_score(y_test,y_pred),4))
                        results.append({"Model":name,
                            "R2 Score": r2,
                            "MAE":      round(mean_absolute_error(y_test,y_pred),4),
                            "RMSE":     round(float(np.sqrt(mean_squared_error(y_test,y_pred))),4)})
                except Exception as me: st.warning(f"⚠️ {name}: {me}")
                prog.progress((i+1)/len(pool), text=f"Trained: {name}")

            prog.empty()
            res_df = pd.DataFrame(results)
            sc = "F1 Score" if task_type=="classification" else "R2 Score"
            if sc in res_df.columns:
                res_df = res_df.sort_values(sc, ascending=False).reset_index(drop=True)
            best_nm = res_df.iloc[0]["Model"] if not res_df.empty else None

            st.session_state.results_df     = res_df
            st.session_state.trained_models = trained_ms
            st.session_state.best_model_name = best_nm
            st.session_state.problem_type    = task_type
            st.session_state.X_test          = X_test
            st.session_state.y_test          = y_test
            st.success(f"✅ Training complete! Best: **{best_nm}**")
            st.rerun()
        except Exception as e:
            st.error(f"❌ Training failed: {e}")
            with st.expander("Debug"): st.code(traceback.format_exc())

    #  Results displayed after training 
    if (st.session_state.results_df is not None and
            st.session_state.problem_type in ["classification","regression"]):
        results_df     = st.session_state.results_df
        trained_models = st.session_state.trained_models
        best_name      = st.session_state.best_model_name
        problem_type   = st.session_state.problem_type
        best_obj       = trained_models.get(best_name)
        exp_rec        = recommend_model_export(best_obj) if best_obj else {}

        st.markdown("---")
        st.markdown(f"""<div class="ai-card">
        🤖 <strong>AI Best Model:</strong> <span class="badge badge-green">{best_name}</span>
        {"— Highest F1 Score (classification)." if problem_type=="classification" else "— Best R² Score (regression)."}
        &nbsp; Export as: <span class="badge badge-purple">{exp_rec.get("format","pickle")}</span>
        — {exp_rec.get("reason","")}
        </div>""", unsafe_allow_html=True)

        st.markdown("#### 📋 Model Leaderboard")
        num_res = results_df.select_dtypes("number").columns.tolist()
        st.dataframe(results_df.style.highlight_max(subset=num_res, color="#7c5cfc33"), use_container_width=True)

        metric_col = "F1 Score" if problem_type=="classification" else "R2 Score"
        if metric_col in results_df.columns:
            fig = px.bar(results_df, x="Model", y=metric_col,
                         color=metric_col, color_continuous_scale=["#5a3fd4","#7c5cfc","#00d4aa"], text=metric_col)
            fig.update_traces(texttemplate='%{text:.4f}', textposition='outside')
            fig.update_layout(**PT, title=f"{metric_col} Comparison")
            st.plotly_chart(fig, use_container_width=True)

        if len(num_res) >= 3:
            fig = go.Figure()
            for _,row in results_df.iterrows():
                vals = [row[m] for m in num_res]
                fig.add_trace(go.Scatterpolar(r=vals+[vals[0]], theta=num_res+[num_res[0]],
                    name=row["Model"], fill='toself', opacity=0.5))
            fig.update_layout(**PT, title="Radar Comparison",
                polar=dict(radialaxis=dict(visible=True,color="#2a2a3a"),bgcolor="rgba(0,0,0,0)"))
            st.plotly_chart(fig, use_container_width=True)

        with st.expander("💾 Download Trained Model (Optional)", expanded=False):
            if best_obj:
                fc1,fc2 = st.columns(2)
                with fc1: fmt_opt = st.selectbox("Format", ["pickle","joblib","onnx"])
                with fc2:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("⬇️ Download Model"):
                        X_s = (st.session_state.X_test.iloc[:5].values
                               if fmt_opt=="onnx" and st.session_state.X_test is not None else None)
                        data, mime = export_model(best_obj, fmt_opt, X_s)
                        if data:
                            ext = {"pickle":"pkl","joblib":"joblib","onnx":"onnx"}.get(fmt_opt,"bin")
                            st.download_button(f"💾 {best_name}.{ext}", data=data,
                                file_name=f"{best_name}.{ext}", mime="application/octet-stream")
                        else: st.error(f"Export failed: {mime}")

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back"): st.session_state.step=6; st.rerun()
    with c2:
        if st.button("Next → Cluster Visualizer ▶"): st.session_state.step=8; st.rerun()
    footer()


# STEP 8 — CLUSTER VISUALIZER (redesigned — no 3D)

elif st.session_state.step == 8:
    sh("📊", "Cluster Visualizer", "Clustering analysis + interactive data graph builder")

    df_c = st.session_state.df_clean
    _, scaled = prepare_clustering_data(df_c)

    if scaled is None:
        st.warning("Not enough numeric data for clustering. Need ≥2 numeric columns & ≥5 rows.")
    else:
        #  SECTION A: Clustering ─
        st.markdown("### 🔵 Clustering Analysis")
        clust_rec = recommend_clustering(df_c)

        # Run all 3 algorithms
        if st.session_state.cluster_results is None:
            with st.spinner("Running KMeans, DBSCAN & Agglomerative..."):
                st.session_state.cluster_results = run_all_clustering(scaled)
        all_clust = st.session_state.cluster_results

        reduced = reduce_to_2d(scaled)
        best_c, best_cs = get_best_clustering(all_clust)

        # Build display df
        num_orig = df_c.select_dtypes(include=[np.number]).dropna()
        n = min(len(reduced), len(num_orig)) if reduced is not None else len(num_orig)
        ddf = num_orig.iloc[:n].reset_index(drop=True).copy()
        if reduced is not None:
            ddf["PC1"] = reduced[:n,0]; ddf["PC2"] = reduced[:n,1]
        num_c = [c for c in ddf.columns if c not in ["PC1","PC2"]]

        # Clustering selector + AI guide
        algo_col, info_col = st.columns([1,2])
        with algo_col:
            st.markdown("#### Choose Algorithm")
            algo_sel = st.selectbox("Clustering Algorithm",
                ["KMeans","DBSCAN","Agglomerative"], key="clust_algo_main")
        with info_col:
            algo_descs = {
                "KMeans": ("Best for large datasets with spherical clusters. Needs K (number of clusters) set manually.", "badge-green"),
                "DBSCAN": ("Best for irregular shapes & noise detection. Does not need K. Slow on large data.", "badge-purple"),
                "Agglomerative": ("Best for small datasets & hierarchical structure. Memory-intensive on large data.", "badge-yellow"),
            }
            desc, badge_cls = algo_descs[algo_sel]
            best_badge = "badge-green" if algo_sel == best_c else "badge-purple"
            st.markdown(f"""<div class="ai-card" style="margin-top:1.5rem">
            🤖 <strong>AI Cluster Guide:</strong><br><br>
            <strong>Recommended:</strong> <span class="badge badge-green">{best_c}</span>
            (Silhouette = <strong>{best_cs:.4f}</strong>
            {"🟢 Excellent" if best_cs>0.5 else "🟡 Moderate" if best_cs>0.25 else "🔴 Poor"})<br><br>
            <strong>Selected ({algo_sel}):</strong> {desc}
            </div>""", unsafe_allow_html=True)

        # Silhouette scores table
        sil_df = pd.DataFrame([{"Algorithm":k,"Silhouette Score":v["score"],"Status":"✅ Best" if k==best_c else ""}
                                for k,v in all_clust.items()]).sort_values("Silhouette Score",ascending=False)
        st.dataframe(sil_df, use_container_width=True)

        # Display clustering scatter (PC1 vs PC2)
        if reduced is not None:
            lbl_sel = [str(all_clust[algo_sel]["labels"][i]) for i in range(n)]
            ddf_plot = ddf.copy(); ddf_plot["Cluster"] = lbl_sel
            fig = px.scatter(ddf_plot, x="PC1", y="PC2", color="Cluster",
                color_discrete_sequence=COLORS, symbol="Cluster",
                title=f"{algo_sel} — Cluster Map (PCA 2D)",
                labels={"PC1":"Principal Component 1","PC2":"Principal Component 2"})
            fig.update_traces(marker=dict(size=10, opacity=0.85, line=dict(width=0.5,color='#0a0a0f')))
            fig.update_layout(**PT)
            st.plotly_chart(fig, use_container_width=True)

            # Elbow curve
            with st.expander("📐 KMeans Elbow Curve (Optimal K finder)", expanded=False):
                Ks, inertias = find_optimal_clusters(scaled)
                valid = [(k,i) for k,i in zip(Ks,inertias) if i is not None]
                if valid:
                    kv,iv = zip(*valid)
                    fig2 = go.Figure(go.Scatter(x=list(kv), y=list(iv), mode="lines+markers",
                        marker=dict(color="#7c5cfc",size=9,symbol="diamond"),
                        line=dict(color="#00d4aa",width=2.5)))
                    fig2.update_layout(**PT, title="Elbow Curve — Optimal K",
                                      xaxis_title="K (Clusters)", yaxis_title="Inertia")
                    st.plotly_chart(fig2, use_container_width=True)

        #  SECTION B: Interactive Graph Builder 
        st.markdown("---")
        st.markdown("### 🎨 Interactive Graph Builder")
        st.markdown("""<div class="ai-card" style="font-size:.85rem">
        📊 <strong>Build any chart from your data.</strong> Select chart type, then configure X/Y axes.
        Compare multiple columns, explore distributions, and spot patterns.
        </div>""", unsafe_allow_html=True)

        all_df_cols = df_c.columns.tolist()
        num_df_cols = df_c.select_dtypes(include=[np.number]).columns.tolist()
        cat_df_cols = df_c.select_dtypes(include=["object","category"]).columns.tolist()

        CHART_TYPES = ["Scatter Plot","Bar Chart","Histogram","Box Plot","Pie Chart",
                       "Violin Plot","Line Chart","Area Chart","Density (KDE)","ECDF Plot",
                       "Regression Plot","Heatmap (Correlation)","Compare Plot"]

        gc1, gc2 = st.columns(2)
        with gc1: chart_type = st.selectbox("📈 Chart Type", CHART_TYPES, key="gb_chart")
        with gc2:
            color_by = st.selectbox("🎨 Color By (optional)", ["None"]+all_df_cols, key="gb_color")
            color_col = None if color_by == "None" else color_by

        # Axis selectors — context aware
        def ax_select(label, options, key, exclude=None):
            opts = [c for c in options if c != exclude] if exclude else options
            if not opts: opts = options
            return st.selectbox(label, opts, key=key)

        fig = None

        if chart_type == "Scatter Plot":
            a1,a2 = st.columns(2)
            with a1: xc = ax_select("X Axis", num_df_cols or all_df_cols, "sc_x")
            with a2: yc = ax_select("Y Axis", num_df_cols or all_df_cols, "sc_y", exclude=xc)
            fig = px.scatter(df_c, x=xc, y=yc, color=color_col, color_discrete_sequence=COLORS,
                title=f"Scatter: {xc} vs {yc}", trendline="ols" if len(df_c)>5 else None)
            fig.update_traces(marker=dict(size=9, opacity=0.8))

        elif chart_type == "Bar Chart":
            a1,a2 = st.columns(2)
            with a1: xc = ax_select("X (Category)", all_df_cols, "bar_x")
            with a2: yc = ax_select("Y (Value)", num_df_cols or all_df_cols, "bar_y", exclude=xc)
            barmode = st.radio("Bar mode", ["group","stack","overlay"], horizontal=True, key="bar_mode")
            fig = px.bar(df_c, x=xc, y=yc, color=color_col or xc,
                color_discrete_sequence=COLORS, barmode=barmode, title=f"Bar: {xc} vs {yc}")

        elif chart_type == "Histogram":
            a1,a2 = st.columns(2)
            with a1: xc = ax_select("Column", num_df_cols or all_df_cols, "hist_x")
            with a2: nbins = st.slider("Bins", 5, 100, 30, key="hist_bins")
            fig = px.histogram(df_c, x=xc, nbins=nbins, color=color_col,
                color_discrete_sequence=COLORS, title=f"Histogram: {xc}", marginal="box")

        elif chart_type == "Box Plot":
            a1,a2 = st.columns(2)
            with a1: yc = ax_select("Y (values)", num_df_cols or all_df_cols, "box_y")
            with a2: xc = ax_select("X (groups)", ["None"]+all_df_cols, "box_x")
            xc = None if xc == "None" else xc
            fig = px.box(df_c, x=xc, y=yc, color=color_col or xc,
                color_discrete_sequence=COLORS, title=f"Box Plot: {yc}", points="all")

        elif chart_type == "Pie Chart":
            a1,a2 = st.columns(2)
            with a1: names_c = ax_select("Labels", all_df_cols, "pie_names")
            with a2:
                val_opts = ["Count"]+num_df_cols
                val_c = st.selectbox("Values", val_opts, key="pie_vals")
            if val_c == "Count":
                pie_data = df_c[names_c].value_counts().reset_index()
                pie_data.columns = [names_c, "Count"]
                fig = px.pie(pie_data, names=names_c, values="Count",
                    color_discrete_sequence=COLORS, title=f"Pie: {names_c} distribution", hole=0.35)
            else:
                fig = px.pie(df_c, names=names_c, values=val_c,
                    color_discrete_sequence=COLORS, title=f"Pie: {val_c} by {names_c}", hole=0.35)
            if fig: fig.update_traces(textposition='inside', textinfo='percent+label')

        elif chart_type == "Violin Plot":
            a1,a2 = st.columns(2)
            with a1: yc = ax_select("Y (values)", num_df_cols or all_df_cols, "vio_y")
            with a2: xc = ax_select("X (groups)", ["None"]+all_df_cols, "vio_x")
            xc = None if xc == "None" else xc
            fig = px.violin(df_c, x=xc, y=yc, color=color_col or xc,
                color_discrete_sequence=COLORS, box=True, points="all",
                title=f"Violin: {yc}")

        elif chart_type == "Line Chart":
            a1,a2 = st.columns(2)
            with a1: xc = ax_select("X Axis", all_df_cols, "line_x")
            with a2: yc = ax_select("Y Axis", num_df_cols or all_df_cols, "line_y", exclude=xc)
            fig = px.line(df_c.sort_values(xc), x=xc, y=yc, color=color_col,
                color_discrete_sequence=COLORS, title=f"Line: {xc} vs {yc}", markers=True)

        elif chart_type == "Area Chart":
            a1,a2 = st.columns(2)
            with a1: xc = ax_select("X Axis", all_df_cols, "area_x")
            with a2: yc = ax_select("Y Axis", num_df_cols or all_df_cols, "area_y", exclude=xc)
            fig = px.area(df_c.sort_values(xc), x=xc, y=yc, color=color_col,
                color_discrete_sequence=COLORS, title=f"Area: {xc} vs {yc}")

        elif chart_type == "Density (KDE)":
            xc = ax_select("Column", num_df_cols or all_df_cols, "kde_x")
            fig = px.histogram(df_c, x=xc, color=color_col, marginal="rug",
                histnorm="density", color_discrete_sequence=COLORS,
                title=f"KDE Density: {xc}")

        elif chart_type == "ECDF Plot":
            xc = ax_select("Column", num_df_cols or all_df_cols, "ecdf_x")
            fig = px.ecdf(df_c, x=xc, color=color_col,
                color_discrete_sequence=COLORS, title=f"ECDF: {xc}")

        elif chart_type == "Regression Plot":
            a1,a2 = st.columns(2)
            with a1: xc = ax_select("X Axis", num_df_cols or all_df_cols, "reg_x")
            with a2: yc = ax_select("Y Axis", num_df_cols or all_df_cols, "reg_y", exclude=xc)
            fig = px.scatter(df_c, x=xc, y=yc, color=color_col,
                color_discrete_sequence=COLORS, trendline="ols",
                title=f"Regression: {xc} vs {yc}")

        elif chart_type == "Heatmap (Correlation)":
            num_only = df_c.select_dtypes(include=[np.number])
            if len(num_only.columns) < 2:
                st.info("Need ≥2 numeric columns for correlation heatmap.")
            else:
                corr = num_only.corr()
                fig = go.Figure(go.Heatmap(z=corr.values, x=corr.columns, y=corr.index,
                    colorscale=[[0,"#ff6b6b"],[.5,"#1a1a26"],[1,"#7c5cfc"]],
                    text=np.round(corr.values,2), texttemplate="%{text}",
                    colorbar=dict(title="Correlation")))
                fig.update_layout(**PT, title="Correlation Heatmap")

        elif chart_type == "Compare Plot":
            st.markdown("##### Compare multiple columns side by side")
            comp_cols = st.multiselect("Select columns to compare", num_df_cols, default=num_df_cols[:3] if len(num_df_cols)>=3 else num_df_cols, key="comp_cols")
            comp_type = st.radio("Compare as", ["Box","Violin","Bar (mean)"], horizontal=True, key="comp_type")
            if comp_cols:
                df_melt = df_c[comp_cols].melt(var_name="Column", value_name="Value")
                if comp_type == "Box":
                    fig = px.box(df_melt, x="Column", y="Value", color="Column",
                        color_discrete_sequence=COLORS, title="Compare: Box Plot", points="all")
                elif comp_type == "Violin":
                    fig = px.violin(df_melt, x="Column", y="Value", color="Column",
                        color_discrete_sequence=COLORS, title="Compare: Violin Plot", box=True)
                else:
                    mean_df = df_c[comp_cols].mean().reset_index(); mean_df.columns=["Column","Mean"]
                    fig = px.bar(mean_df, x="Column", y="Mean", color="Column",
                        color_discrete_sequence=COLORS, title="Compare: Mean Values", text="Mean")
                    if fig: fig.update_traces(texttemplate='%{text:.3f}', textposition='outside')

        if fig:
            fig.update_layout(**PT)
            st.plotly_chart(fig, use_container_width=True)

    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back to Training"): st.session_state.step=7; st.rerun()
    with c2:
        if st.button("Next → Report & Export ▶"): st.session_state.step=9; st.rerun()
    footer()


# STEP 9 — REPORT & EXPORT (professional Excel + PDF + Word)

elif st.session_state.step == 9:
    sh("📑", "Report & Export", "Download your complete ML session report — Excel, PDF & Word")

    df_clean     = st.session_state.df_clean
    results_df   = st.session_state.results_df
    best_name    = st.session_state.best_model_name
    problem_type = st.session_state.problem_type
    cluster_res  = st.session_state.cluster_results
    ai_report    = st.session_state.ai_report or {}

    st.markdown("""<div class="ai-card">
    🤖 <strong>AI Report Guide:</strong>
    Select sections to include below. All 3 formats contain the same data with professional layouts.
    <strong>Excel</strong> = data analysis with colored tables & embedded charts.
    <strong>PDF</strong> = print-ready professional report with color sections.
    <strong>Word</strong> = editable colored report for further customization.
    </div>""", unsafe_allow_html=True)

    st.markdown("#### ✅ Select Sections to Include")
    OPTS = {
        "📊 Dataset Overview":               df_clean is not None,
        "🔍 Column Profiling Summary":       df_clean is not None,
        "🧹 Missing Value Handling Log":     st.session_state.missing_applied,
        "♻️ Duplicate Removal Log":          st.session_state.dup_removed,
        "⚙️ Preprocessing Steps":            len(st.session_state.preprocessing_log) > 0,
        "📋 Clean Dataset (first 500 rows)": df_clean is not None,
        "🤖 Model Training Results":         results_df is not None,
        "🔵 Clustering Silhouette Scores":   cluster_res is not None,
        "💡 AI Recommendations":             True,
    }
    selected = {}
    cols_c = st.columns(2)
    for i,(lbl,avail) in enumerate(OPTS.items()):
        with cols_c[i%2]:
            if avail:
                selected[lbl] = st.checkbox(lbl, value=True, key=f"chk_{i}")
            else:
                st.checkbox(lbl, value=False, disabled=True, key=f"chkd_{i}",
                            help="Step not completed")
                selected[lbl] = False

    st.markdown("---")

    #  Equal-height format cards 
    CARD_H = "170px"
    fc1, fc2, fc3 = st.columns(3)

    for col, fmt_icon, fmt_title, fmt_desc in [
        (fc1, "📊", "Excel Report", "Multi-sheet · Colored tables · Charts · AI Guide"),
        (fc2, "📄", "PDF Report",   "Professional · Color sections · Print-ready · AI insights"),
        (fc3, "📝", "Word Report",  "Editable .docx · Colored tables · Section headers"),
    ]:
        with col:
            st.markdown(f"""<div class="metric-card" style="min-height:{CARD_H};margin-bottom:1rem">
            <div style="font-size:2.5rem;margin-bottom:.3rem">{fmt_icon}</div>
            <div style="font-weight:700;color:var(--accent2);font-size:1.05rem">{fmt_title}</div>
            <div style="font-size:.76rem;color:var(--muted);margin-top:.4rem;line-height:1.5">{fmt_desc}</div>
            </div>""", unsafe_allow_html=True)

    fc1, fc2, fc3 = st.columns(3)

    # ════ EXCEL ══════════════════════════════════════════════
    with fc1:
        if st.button("⬇️ Download Excel (.xlsx)", use_container_width=True):
            try:
                import datetime
                from openpyxl import load_workbook
                from openpyxl.styles import (PatternFill, Font, Alignment, Border, Side, GradientFill)
                from openpyxl.utils import get_column_letter
                from openpyxl.chart import BarChart, Reference
                from openpyxl.chart.series import SeriesLabel

                buf = io.BytesIO()
                with pd.ExcelWriter(buf, engine="openpyxl") as w:
                    def ws(df_t, name):
                        df_t.to_excel(w, sheet_name=name[:31], index=False)

                    # Always write cover
                    cover_data = [["ML Studio Pro — Session Report", ""],
                                  ["Generated", datetime.datetime.now().strftime("%Y-%m-%d %H:%M")],
                                  ["Rows", f"{len(df_clean):,}" if df_clean is not None else "N/A"],
                                  ["Columns", str(len(df_clean.columns)) if df_clean is not None else "N/A"],
                                  ["Best Model", best_name or "N/A"],
                                  ["Task Type", problem_type or "N/A"],
                                  ["Training Dataset", st.session_state.get("training_dataset", "preprocessed").title()]]
                    ws(pd.DataFrame(cover_data, columns=["Field","Value"]), "📋 Cover")

                    if selected.get("📊 Dataset Overview") and df_clean is not None:
                        info = get_basic_info(df_clean)
                        ws(pd.DataFrame([
                            ["Total Rows", f"{info['rows']:,}"],
                            ["Total Columns", str(info["columns"])],
                            ["Missing (post-clean)", str(int(df_clean.isna().sum().sum()))],
                            ["Duplicates (original)", str(info["duplicate_rows"])],
                            ["Best Model", best_name or "N/A"],
                            ["Task Type", problem_type or "N/A"],
                        ], columns=["Metric","Value"]), "📊 Overview")

                    if selected.get("🔍 Column Profiling Summary") and df_clean is not None:
                        ws(get_column_summary(df_clean), "🔍 Column Summary")
                        ns = get_numeric_stats(df_clean)
                        if not ns.empty: ws(ns, "📈 Numeric Stats")

                    if selected.get("⚙️ Preprocessing Steps") and st.session_state.preprocessing_log:
                        ws(pd.DataFrame({"Step": range(1, len(st.session_state.preprocessing_log)+1),
                                         "Action": st.session_state.preprocessing_log}), "⚙️ Preprocessing")

                    if selected.get("📋 Clean Dataset (first 500 rows)") and df_clean is not None:
                        ws(df_clean.head(500), "📋 Clean Data")

                    if selected.get("🤖 Model Training Results") and results_df is not None:
                        ws(results_df, "🤖 Model Results")

                    if selected.get("🔵 Clustering Silhouette Scores") and cluster_res:
                        ws(pd.DataFrame([{"Algorithm":k,"Silhouette Score":v["score"]}
                                          for k,v in cluster_res.items()]), "🔵 Clustering")

                    if selected.get("💡 AI Recommendations"):
                        rows = []
                        if ai_report.get("problem_type"): rows.append(["Recommended Task", ai_report["problem_type"]])
                        for m in ai_report.get("recommended_models",[]): rows.append(["Recommended Model", m])
                        for f in ai_report.get("important_features",[]): rows.append(["Top Feature", str(f)])
                        cl = ai_report.get("clustering",{})
                        if cl.get("recommended"): rows.append(["Best Clustering", f"{cl['recommended']} — {cl.get('reason','')}"])
                        if rows: ws(pd.DataFrame(rows, columns=["Category","Value"]), "💡 AI Guide")

                # Apply ultra-pro styling
                buf.seek(0)
                wb = load_workbook(buf)

                # Color palette
                C_NAVY  = "1E3A5F"
                C_TEAL  = "00D4AA"
                C_PURP  = "7C5CFC"
                C_GOLD  = "FFB347"
                C_ALT1  = "EEF2FF"
                C_ALT2  = "E8FFF9"
                C_ALT3  = "FFF8EE"

                # Sheet-specific accent colors
                SHEET_COLORS = {
                    "📋 Cover":       (C_NAVY, C_ALT1, "1E3A5F"),
                    "📊 Overview":    (C_NAVY, C_ALT1, "1E3A5F"),
                    "🔍 Column Summary": (C_PURP, C_ALT1, "7C5CFC"),
                    "📈 Numeric Stats":  (C_TEAL, C_ALT2, "00608A"),
                    "⚙️ Preprocessing": (C_PURP, C_ALT1, "7C5CFC"),
                    "📋 Clean Data":    (C_NAVY, C_ALT1, "1E3A5F"),
                    "🤖 Model Results": (C_NAVY, C_ALT1, "1E3A5F"),
                    "🔵 Clustering":    (C_TEAL, C_ALT2, "00608A"),
                    "💡 AI Guide":      (C_GOLD, C_ALT3, "B8860B"),
                }

                thin = Side(style="thin", color="C0C0D8")
                BDR = Border(left=thin, right=thin, top=thin, bottom=thin)

                for ws_obj in wb.worksheets:
                    sname = ws_obj.title
                    hdr_c, alt_c, tab_c = SHEET_COLORS.get(sname, (C_NAVY, C_ALT1, "1E3A5F"))
                    HFill  = PatternFill("solid", fgColor=hdr_c)
                    AFill  = PatternFill("solid", fgColor=alt_c)
                    WHFill = PatternFill("solid", fgColor="FFFFFF")
                    WFont  = Font(color="FFFFFF", bold=True, name="Calibri", size=10)
                    HFont  = Font(color=hdr_c, bold=True, name="Calibri", size=9)
                    DFont  = Font(color="1E1E2E", name="Calibri", size=9)

                    # Auto-width
                    for col_cells in ws_obj.columns:
                        max_l = max((len(str(c.value or "")) for c in col_cells), default=8)
                        ws_obj.column_dimensions[get_column_letter(col_cells[0].column)].width = min(max_l + 4, 50)

                    # Style rows
                    for ri, row in enumerate(ws_obj.iter_rows()):
                        for cell in row:
                            cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
                            if ri == 0:
                                cell.fill = HFill; cell.font = WFont; cell.border = BDR
                                ws_obj.row_dimensions[ri+1].height = 24
                            elif ri % 2 == 1:
                                cell.fill = AFill; cell.font = DFont; cell.border = BDR
                            else:
                                cell.fill = WHFill; cell.font = DFont; cell.border = BDR

                    ws_obj.freeze_panes = "A2"
                    ws_obj.auto_filter.ref = ws_obj.dimensions
                    ws_obj.sheet_properties.tabColor = tab_c

                # Conditional formatting on Model Results
                if "🤖 Model Results" in wb.sheetnames and results_df is not None:
                    from openpyxl.formatting.rule import ColorScaleRule
                    ws_m = wb["🤖 Model Results"]
                    n_rows = len(results_df) + 1
                    n_cols = len(results_df.columns)
                    for ci in range(2, n_cols + 1):
                        col_letter = get_column_letter(ci)
                        rng = f"{col_letter}2:{col_letter}{n_rows}"
                        ws_m.conditional_formatting.add(rng, ColorScaleRule(
                            start_type="min", start_color="FF6B6B",
                            mid_type="percentile", mid_value=50, mid_color="FFEB84",
                            end_type="max", end_color="00D4AA"
                        ))
                    # Bar chart
                    try:
                        chart = BarChart()
                        chart.type = "col"; chart.grouping = "clustered"
                        chart.title = "Model Performance Comparison"
                        chart.style = 10
                        data_ref = Reference(ws_m, min_col=2, max_col=n_cols, min_row=1, max_row=n_rows)
                        cats_ref = Reference(ws_m, min_col=1, min_row=2, max_row=n_rows)
                        chart.add_data(data_ref, titles_from_data=True)
                        chart.set_categories(cats_ref)
                        chart.width = 22; chart.height = 14
                        ws_m.add_chart(chart, f"A{n_rows+3}")
                    except Exception: pass

                styled = io.BytesIO(); wb.save(styled); styled.seek(0)
                st.download_button("💾 ML_Studio_Report.xlsx", data=styled.getvalue(),
                    file_name="ML_Studio_Report.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                st.success("✅ Excel ready!")
            except Exception as e:
                st.error(f"❌ {e}")
                with st.expander("Debug"): st.code(traceback.format_exc())

    # ════ PDF ════════════════════════════════════════════════
    with fc2:
        if st.button("⬇️ Download PDF (.pdf)", use_container_width=True):
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib import colors
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import cm
                from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                                Table, TableStyle, HRFlowable, KeepTogether,
                                                PageBreak)
                from reportlab.lib.enums import TA_CENTER, TA_LEFT
                from reportlab.graphics.shapes import Drawing, Rect, String
                from reportlab.graphics.charts.barcharts import VerticalBarChart
                from reportlab.graphics import renderPDF

                W, H = A4
                NAVY  = colors.HexColor("#1E3A5F")
                PRP   = colors.HexColor("#7c5cfc")
                TEL   = colors.HexColor("#00d4aa")
                YEL   = colors.HexColor("#ffb347")
                RED   = colors.HexColor("#ff6b6b")
                GRY   = colors.HexColor("#666680")
                BGLT  = colors.HexColor("#EEF2FF")
                BGTL  = colors.HexColor("#E8FFF9")
                WHT   = colors.white
                DRK   = colors.HexColor("#1E1E2E")

                pdf_buf = io.BytesIO()
                doc = SimpleDocTemplate(pdf_buf, pagesize=A4,
                    leftMargin=1.8*cm, rightMargin=1.8*cm, topMargin=2*cm, bottomMargin=2*cm)
                S = getSampleStyleSheet()

                sT   = ParagraphStyle("T",  textColor=NAVY,  fontSize=28, fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=4)
                sSB  = ParagraphStyle("SB", textColor=TEL,   fontSize=11, fontName="Helvetica",      alignment=TA_CENTER, spaceAfter=16)
                sH2  = ParagraphStyle("H2", textColor=WHT,   fontSize=12, fontName="Helvetica-Bold", spaceBefore=14, spaceAfter=0,
                                      backColor=NAVY, leftIndent=0, rightIndent=0, borderPad=7)
                sBD  = ParagraphStyle("BD", fontSize=9,  fontName="Helvetica", textColor=DRK, spaceAfter=3, leftIndent=10)
                sMT  = ParagraphStyle("MT", fontSize=8,  fontName="Helvetica", textColor=GRY, spaceAfter=2, leftIndent=10)

                def sec_hdr(txt, color=NAVY):
                    bg = ParagraphStyle("sh", textColor=WHT, fontSize=12, fontName="Helvetica-Bold",
                                        spaceBefore=14, spaceAfter=4, backColor=color,
                                        leftIndent=0, rightIndent=0, borderPad=8)
                    return Paragraph(f"  {txt}", bg)

                def make_table(data, cw=None, hdr=NAVY, alt=BGLT):
                    t = Table(data, colWidths=cw, repeatRows=1)
                    t.setStyle(TableStyle([
                        ("BACKGROUND",    (0,0), (-1,0), hdr),
                        ("TEXTCOLOR",     (0,0), (-1,0), WHT),
                        ("FONTNAME",      (0,0), (-1,0), "Helvetica-Bold"),
                        ("FONTSIZE",      (0,0), (-1,-1), 8.5),
                        ("ALIGN",         (0,0), (-1,-1), "CENTER"),
                        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
                        ("PADDING",       (0,0), (-1,-1), 5),
                        ("GRID",          (0,0), (-1,-1), 0.4, colors.HexColor("#B0B8E0")),
                        ("ROWBACKGROUNDS",(0,1), (-1,-1), [WHT, alt]),
                        ("BOTTOMPADDING", (0,0), (-1,0), 7),
                    ]))
                    return t

                def bar_chart_pdf(labels, values, title="", color=PRP):
                    drawing = Drawing(460, 220)
                    bc = VerticalBarChart()
                    bc.x = 50; bc.y = 60; bc.width = 380; bc.height = 130
                    bc.data = [values]
                    bc.bars[0].fillColor = color
                    # Truncate long labels so they don't overlap
                    bc.categoryAxis.categoryNames = [str(L)[:15]+".." if len(str(L))>16 else str(L) for L in labels]
                    bc.categoryAxis.labels.angle = 45
                    bc.categoryAxis.labels.dx = 0
                    bc.categoryAxis.labels.dy = -10
                    bc.categoryAxis.labels.fontSize = 7
                    bc.valueAxis.labels.fontSize = 7
                    
                    # Adjust y-axis range
                    v_min = min(values) if values else 0
                    v_max = max(values) if values else 1
                    bc.valueAxis.valueMin = v_min * 1.1 if v_min < 0 else 0
                    bc.valueAxis.valueMax = v_max * 1.2
                    
                    bc.strokeColor = colors.HexColor("#B0B8E0")
                    
                    # Value labels on top of bars
                    bc.barLabelFormat = "%.2f"
                    bc.barLabels.nudge = 5
                    bc.barLabels.fontSize = 6
                    bc.barLabels.boxAnchor = 's'
                    
                    drawing.add(bc)
                    drawing.add(String(230, 205, title, textAnchor='middle',
                                       fontSize=10, fillColor=NAVY, fontName="Helvetica-Bold"))
                    return drawing


                import datetime as _dt

                # Page-number canvas maker
                class _PageNum:
                    def __init__(self, doc):
                        self.doc = doc
                    def afterPage(self):
                        pass
                def _add_page_num(canvas, doc):
                    canvas.saveState()
                    canvas.setFont("Helvetica", 8)
                    canvas.setFillColor(GRY)
                    canvas.drawCentredString(W/2, 1.2*cm, f"ML Studio Pro — Page {doc.page}")
                    canvas.restoreState()

                # Cover page elements
                _now = _dt.datetime.now().strftime("%B %d, %Y  %H:%M")
                _rows = len(df_clean) if df_clean is not None else 0
                _cols = len(df_clean.columns) if df_clean is not None else 0

                # Large cover block
                _cover_table_data = [["ML Studio Pro", "Session Report"],
                                     ["Generated", _now],
                                     ["Dataset", f"{_rows:,} rows × {_cols} columns"],
                                     ["Best Model", best_name or "N/A"],
                                     ["Task Type", (problem_type or "N/A").title()]]
                _ct = Table(_cover_table_data, colWidths=[6*cm, 9*cm])
                _ct.setStyle(TableStyle([
                    ("BACKGROUND", (0,0), (-1,0), NAVY),
                    ("TEXTCOLOR",  (0,0), (-1,0), TEL),
                    ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
                    ("FONTSIZE",   (0,0), (-1,0), 16),
                    ("BACKGROUND", (0,1), (-1,-1), colors.HexColor("#EEF2FF")),
                    ("TEXTCOLOR",  (0,1), (-1,-1), colors.HexColor("#1E1E2E")),
                    ("FONTNAME",   (0,1), (0,-1), "Helvetica-Bold"),
                    ("FONTNAME",   (1,1), (1,-1), "Helvetica"),
                    ("FONTSIZE",   (0,1), (-1,-1), 10),
                    ("ALIGN",      (0,0), (-1,-1), "LEFT"),
                    ("PADDING",    (0,0), (-1,-1), 10),
                    ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#C0C0D8")),
                    ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.HexColor("#EEF2FF"), colors.white]),
                ]))

                story = [
                    Spacer(1, 1.5*cm),
                    Paragraph("⚡  ML Studio Pro", sT),
                    Paragraph("Machine Learning Session Report", sSB),
                    HRFlowable(width="100%", color=NAVY, thickness=4),
                    Spacer(1, 0.8*cm),
                    _ct,
                    Spacer(1, 0.5*cm),
                    HRFlowable(width="100%", color=TEL, thickness=1),
                    Spacer(1, 0.5*cm),
                ]

                if selected.get("📊 Dataset Overview") and df_clean is not None:
                    info = get_basic_info(df_clean)
                    story += [sec_hdr("📊  Dataset Overview"), Spacer(1,4)]
                    data = [["Metric","Value"],
                            ["Total Rows", f"{info['rows']:,}"],
                            ["Total Columns", str(info["columns"])],
                            ["Missing (post-clean)", str(int(df_clean.isna().sum().sum()))],
                            ["Best Model", best_name or "N/A"],
                            ["Task Type", problem_type or "N/A"]]
                    story += [make_table(data, cw=[8.5*cm,6.5*cm]), Spacer(1,10)]

                if selected.get("🔍 Column Profiling Summary") and df_clean is not None:
                    story += [sec_hdr("🔍  Column Profiling"), Spacer(1,4)]
                    cs = get_column_summary(df_clean)
                    hdr_r = list(cs.columns)
                    rows_r = [[str(v) for v in r] for _,r in cs.head(15).iterrows()]
                    cw_r = [(W-3.6*cm)/len(hdr_r)]*len(hdr_r)
                    story += [make_table([hdr_r]+rows_r, cw=cw_r, alt=BGTL, hdr=PRP), Spacer(1,10)]

                if selected.get("⚙️ Preprocessing Steps") and st.session_state.preprocessing_log:
                    story += [sec_hdr("⚙️  Preprocessing Steps"), Spacer(1,4)]
                    pdata = [["#","Action"]] + [[str(i+1), step] for i,step in enumerate(st.session_state.preprocessing_log)]
                    story += [make_table(pdata, cw=[1.5*cm, 13*cm], hdr=PRP), Spacer(1,10)]

                if selected.get("🤖 Model Training Results") and results_df is not None:
                    story += [sec_hdr("🤖  Model Training Results"), Spacer(1,4)]
                    hdr_r = list(results_df.columns)
                    rows_r = [[str(round(v,4) if isinstance(v,float) else v) for v in r] for _,r in results_df.iterrows()]
                    cw_r = [(W-3.6*cm)/len(hdr_r)]*len(hdr_r)
                    story += [make_table([hdr_r]+rows_r, cw=cw_r), Spacer(1,6)]
                    story.append(Paragraph(f"  🏆  Best Model: {best_name}", sBD))
                    # Add bar chart
                    num_res = results_df.select_dtypes("number").columns.tolist()
                    if num_res and len(results_df) > 0:
                        metric_c = "F1 Score" if "F1 Score" in num_res else (num_res[0] if num_res else None)
                        if metric_c:
                            story += [Spacer(1,8),
                                      bar_chart_pdf(results_df["Model"].tolist(),
                                                    [float(v) for v in results_df[metric_c].tolist()],
                                                    title=f"{metric_c} by Model"),
                                      Spacer(1,10)]

                if selected.get("🔵 Clustering Silhouette Scores") and cluster_res:
                    story += [sec_hdr("🔵  Clustering Results", color=colors.HexColor("#00608A")), Spacer(1,4)]
                    cdata = [["Algorithm","Silhouette Score"]] + [[k, str(v["score"])] for k,v in cluster_res.items()]
                    story += [make_table(cdata, cw=[8.5*cm,6.5*cm], hdr=colors.HexColor("#00608A"), alt=BGTL),
                              Spacer(1,8)]
                    # Cluster bar chart
                    cl_labels = list(cluster_res.keys())
                    cl_scores = [float(v["score"]) for v in cluster_res.values()]
                    story += [bar_chart_pdf(cl_labels, cl_scores, "Silhouette Scores by Algorithm"), Spacer(1,10)]

                if selected.get("💡 AI Recommendations"):
                    story += [sec_hdr("💡  AI Recommendations", color=colors.HexColor("#B8860B")), Spacer(1,6)]
                    pts = []
                    if ai_report.get("problem_type"): pts.append(f"🎯  Problem Type: <b>{ai_report['problem_type']}</b>")
                    if ai_report.get("recommended_models"): pts.append("🧠  Models: " + ", ".join(ai_report["recommended_models"][:5]))
                    if ai_report.get("important_features"): pts.append("⭐  Top Features: " + ", ".join([str(f) for f in ai_report["important_features"][:5]]))
                    cl = ai_report.get("clustering",{})
                    if cl.get("recommended"): pts.append(f"🔵  Best Clustering: {cl['recommended']} — {cl.get('reason','')}")
                    for pt in pts: story.append(Paragraph(f"• {pt}", sBD))

                story += [Spacer(1,1*cm), HRFlowable(width="100%", color=GRY, thickness=1), Spacer(1,6),
                          Paragraph("Generated by ⚡ ML Studio Pro v2.0  |  Advanced AutoML Platform", sSB)]

                doc.build(story, onFirstPage=_add_page_num, onLaterPages=_add_page_num)
                pdf_buf.seek(0)
                st.download_button("💾 ML_Studio_Report.pdf", data=pdf_buf.getvalue(),
                    file_name="ML_Studio_Report.pdf", mime="application/pdf")
                st.success("✅ PDF ready!")
            except Exception as e:
                st.error(f"❌ PDF error: {e}")
                with st.expander("Debug"): st.code(traceback.format_exc())

    # ════ WORD ═══════════════════════════════════════════════
    with fc3:
        if st.button("⬇️ Download Word (.docx)", use_container_width=True):
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor, Cm, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement

                dw = Document()
                for sec in dw.sections:
                    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
                    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

                def set_cell_bg(cell, hex_c):
                    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
                    shd.set(qn('w:fill'), hex_c); tcPr.append(shd)

                def add_para_border(para, hex_c="1E3A5F", sz=24):
                    """Add a bottom border to a paragraph."""
                    from docx.oxml import OxmlElement
                    pPr = para._p.get_or_add_pPr()
                    pBdr = OxmlElement('w:pBdr')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), str(sz))
                    bottom.set(qn('w:color'), hex_c)
                    pBdr.append(bottom)
                    pPr.append(pBdr)

                import datetime as _dt2
                _now_w = _dt2.datetime.now().strftime("%B %d, %Y  %H:%M")

                #  Cover page 
                # Big gradient-like header using a 1-row table
                cover_tb = dw.add_table(rows=1, cols=1)
                cover_tb.style = "Table Grid"
                cover_cell = cover_tb.rows[0].cells[0]
                set_cell_bg(cover_cell, "1E3A5F")
                cover_cell.width = Cm(16)
                cp1 = cover_cell.paragraphs[0]
                cp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr1 = cp1.add_run("⚡  ML Studio Pro")
                cr1.font.size = Pt(28); cr1.font.bold = True
                cr1.font.color.rgb = RGBColor(0x00,0xD4,0xAA)
                cp2 = cover_cell.add_paragraph()
                cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr2 = cp2.add_run("Machine Learning Session Report")
                cr2.font.size = Pt(13); cr2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

                dw.add_paragraph()

                # KPI info table on cover
                _cv_rows = len(df_clean) if df_clean is not None else 0
                _cv_cols = len(df_clean.columns) if df_clean is not None else 0
                kpi_data = [["Field","Value"],
                             ["Generated", _now_w],
                             ["Dataset", f"{_cv_rows:,} rows × {_cv_cols} columns"],
                             ["Best Model", best_name or "N/A"],
                             ["Task Type", (problem_type or "N/A").title()],
                             ["Training Dataset", st.session_state.training_dataset.title()]]
                kpi_tb = dw.add_table(rows=len(kpi_data), cols=2)
                kpi_tb.style = "Table Grid"
                for ri, row_d in enumerate(kpi_data):
                    cells = kpi_tb.rows[ri].cells
                    for ci, val in enumerate(row_d):
                        cells[ci].text = str(val)
                        bg = "1E3A5F" if ri == 0 else ("EEF2FF" if ri % 2 == 1 else "FFFFFF")
                        set_cell_bg(cells[ci], bg)
                        for run in cells[ci].paragraphs[0].runs:
                            run.font.size = Pt(10)
                            if ri == 0:
                                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                                run.font.bold = True
                            elif ci == 0:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0x1E,0x3A,0x5F)
                        cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                dw.add_paragraph()
                # Page break after cover
                import docx
                from docx.oxml import OxmlElement as _oe
                _pb = _oe('w:p')
                _ppr = _oe('w:pPr')
                _pbr = _oe('w:pageBreak')
                dw.add_paragraph().runs  # placeholder
                run_pb = dw.paragraphs[-1].add_run()
                run_pb.add_break(docx.enum.text.WD_BREAK.PAGE) if hasattr(docx, 'enum') else None
                try:
                    from docx.enum.text import WD_BREAK
                    dw.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
                except Exception: pass

                def add_sec(icon, title, rgb=(0x1E,0x3A,0x5F)):
                    h = dw.add_heading("", level=1); h.clear()
                    run = h.add_run(f"{icon}  {title}"); run.font.size=Pt(14); run.font.bold=True
                    run.font.color.rgb = RGBColor(*rgb)

                def add_table(df_t, hdr_hex="1E3A5F", alt_hex="EEF2FF", max_rows=50):
                    df_t = df_t.head(max_rows)
                    tb = dw.add_table(rows=1+len(df_t), cols=len(df_t.columns))
                    tb.style = "Table Grid"
                    for i,col in enumerate(df_t.columns):
                        cell = tb.rows[0].cells[i]
                        cell.text = str(col); set_cell_bg(cell, hdr_hex)
                        for run in cell.paragraphs[0].runs:
                            run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); run.font.bold=True; run.font.size=Pt(9)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for ri,(_,row) in enumerate(df_t.iterrows()):
                        tr = tb.rows[ri+1]
                        bg = alt_hex if ri%2==0 else "FFFFFF"
                        for i,val in enumerate(row):
                            cell = tr.cells[i]
                            cell.text = str(round(val,4) if isinstance(val,float) else val)
                            set_cell_bg(cell, bg)
                            for run in cell.paragraphs[0].runs:
                                run.font.size=Pt(9); run.font.color.rgb=RGBColor(0x1E,0x1E,0x2E)
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    dw.add_paragraph()

                def add_bullet(text, rgb=(0x00,0xD4,0xAA)):
                    p = dw.add_paragraph(); r = p.add_run(f"▸  {text}")
                    r.font.size=Pt(10); r.font.color.rgb=RGBColor(*rgb)

                if selected.get("📊 Dataset Overview") and df_clean is not None:
                    add_sec("📊","Dataset Overview")
                    info = get_basic_info(df_clean)
                    add_table(pd.DataFrame([
                        ["Total Rows",f"{info['rows']:,}"],["Total Columns",str(info["columns"])],
                        ["Missing (post-clean)",str(int(df_clean.isna().sum().sum()))],
                        ["Best Model",best_name or "N/A"],["Task Type",problem_type or "N/A"],
                    ], columns=["Metric","Value"]), hdr_hex="1E3A5F", alt_hex="EEF2FF")

                if selected.get("🔍 Column Profiling Summary") and df_clean is not None:
                    add_sec("🔍","Column Profiling", rgb=(0x00,0x60,0x8A))
                    add_table(get_column_summary(df_clean), hdr_hex="00608A", alt_hex="E8FFF9", max_rows=30)

                if selected.get("⚙️ Preprocessing Steps") and st.session_state.preprocessing_log:
                    add_sec("⚙️","Preprocessing Steps", rgb=(0x7c,0x5c,0xfc))
                    add_table(pd.DataFrame({"Step":range(1,len(st.session_state.preprocessing_log)+1),
                                            "Action":st.session_state.preprocessing_log}),
                              hdr_hex="7C5CFC", alt_hex="F0EEFF", max_rows=100)

                if selected.get("📋 Clean Dataset (first 500 rows)") and df_clean is not None:
                    add_sec("📋","Clean Dataset (first 20 rows)", rgb=(0x5a,0x3f,0xd4))
                    add_table(df_clean.head(20), hdr_hex="5A3FD4", alt_hex="EDE8FF")

                if selected.get("🤖 Model Training Results") and results_df is not None:
                    add_sec("🤖","Model Training Results")
                    add_table(results_df, hdr_hex="1E3A5F", alt_hex="EEF2FF")
                    p = dw.add_paragraph(); r = p.add_run(f"🏆  Best Model: {best_name}")
                    r.font.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x00,0xD4,0xAA)
                    dw.add_paragraph()

                if selected.get("🔵 Clustering Silhouette Scores") and cluster_res:
                    add_sec("🔵","Clustering Results", rgb=(0x00,0x60,0x8A))
                    add_table(pd.DataFrame([{"Algorithm":k,"Silhouette Score":v["score"]} for k,v in cluster_res.items()]),
                              hdr_hex="00608A", alt_hex="E8FFF9")

                if selected.get("💡 AI Recommendations"):
                    add_sec("💡","AI Recommendations", rgb=(0xB8,0x86,0x0B))
                    pts = []
                    if ai_report.get("problem_type"): pts.append(f"Problem Type: {ai_report['problem_type']}")
                    for m in ai_report.get("recommended_models",[])[:5]: pts.append(f"Recommended Model: {m}")
                    for f in ai_report.get("important_features",[])[:5]: pts.append(f"Top Feature: {f}")
                    cl = ai_report.get("clustering",{})
                    if cl.get("recommended"): pts.append(f"Best Clustering: {cl['recommended']} — {cl.get('reason','')}")
                    for pt in pts: add_bullet(pt)

                # Add a native footer with page numbers
                for section in dw.sections:
                    footer = section.footer
                    fp = footer.paragraphs[0]
                    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Add border top to footer
                    add_para_border(fp, hex_c="B0B8E0", sz=12)
                    fr = fp.add_run("Generated by ⚡ ML Studio Pro v2.0  |  Advanced AutoML Platform  |  ")
                    fr.font.color.rgb = RGBColor(0x66,0x66,0x80); fr.font.size=Pt(8)
                    
                    # Page numbers in Word using field codes
                    fldChar1 = OxmlElement('w:fldChar')
                    fldChar1.set(qn('w:fldCharType'), 'begin')
                    instrText = OxmlElement('w:instrText')
                    instrText.set(qn('xml:space'), 'preserve')
                    instrText.text = "PAGE"
                    fldChar2 = OxmlElement('w:fldChar')
                    fldChar2.set(qn('w:fldCharType'), 'end')
                    
                    r_element = fp.add_run()._r
                    r_element.append(fldChar1)
                    r_element.append(instrText)
                    r_element.append(fldChar2)
                    
                    fp.runs[-1].font.color.rgb = RGBColor(0x66,0x66,0x80)
                    fp.runs[-1].font.size=Pt(8)

                wb2 = io.BytesIO(); dw.save(wb2); wb2.seek(0)
                st.download_button("💾 ML_Studio_Report.docx", data=wb2.getvalue(),
                    file_name="ML_Studio_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.success("✅ Word doc ready!")
            except Exception as e:
                st.error(f"❌ Word error: {e}")
                with st.expander("Debug"): st.code(traceback.format_exc())

    #  New Pipeline 
    st.markdown("---")
    st.markdown("""<div class="ai-card" style="text-align:center">
    <div style="font-size:1.5rem">🔄</div>
    <strong>Start a New Pipeline</strong><br>
    <span style="color:var(--muted);font-size:.85rem">Reset everything and begin fresh with a new dataset</span>
    </div>""", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    c1,c2 = st.columns(2)
    with c1:
        if st.button("◀ Back to Visualizer"): st.session_state.step=8; st.rerun()
    with c2:
        if st.button("🔄 Start New Pipeline", use_container_width=True):
            for k in list(st.session_state.keys()): del st.session_state[k]
            for k,v in DEFAULTS.items(): st.session_state[k] = v
            st.rerun()
    footer()
